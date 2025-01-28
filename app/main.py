# Библиотеки
# ---------------------------------------------------------------------------------------------------------------------------------#
from datetime import date, datetime, timedelta
import jwt
import os
from fastapi import FastAPI, HTTPException, Depends, Request, Form, Response, UploadFile
from fastapi.security import OAuth2PasswordBearer
from fastapi.responses import HTMLResponse, JSONResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel
from app.models import get_user_by_username, create_user
from app.utils import check_password, hash_password
from app.database import get_db_connection, close_db_connection
from typing import Optional, List, Tuple, Dict

# ---------------------------------------------------------------------------------------------------------------------------------#


# Константы
# ---------------------------------------------------------------------------------------------------------------------------------#
SECRET_KEY = "your_secret_key"
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 30
# ---------------------------------------------------------------------------------------------------------------------------------#


# Инициализация
# ---------------------------------------------------------------------------------------------------------------------------------#
app = FastAPI()
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="login")
templates = Jinja2Templates(directory="app/templates")
app.mount("/static", StaticFiles(directory="app/static"), name="static")
# ---------------------------------------------------------------------------------------------------------------------------------#


# Функции
# ---------------------------------------------------------------------------------------------------------------------------------#
def create_access_token(data: dict, expires_delta: Optional[timedelta] = None) -> str:
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)

    to_encode = data.copy()
    to_encode.update({"exp": expire})

    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt


def get_current_user(request: Request):
    token = request.headers.get("Authorization")
    if not token:
        raise HTTPException(status_code=401, detail="Not authenticated")

    try:
        token = token.split(" ")[1]  # Убираем "Bearer " из заголовка
        payload = jwt.decode(token, "your_secret_key", algorithms=["HS256"])
        username = payload.get("sub")
        if not username:
            raise HTTPException(status_code=401, detail="Invalid token")

        user = get_user_by_username(username)
        if user is None:
            raise HTTPException(status_code=404, detail="User not found")

        return user
    except jwt.JWTError:
        raise HTTPException(status_code=401, detail="Invalid token")


# Функция для проверки токена и получения текущего пользователя
def get_current_user(request: Request):
    token = request.headers.get("Authorization")
    print("Полученный токен:", token)  # Логирование токена

    if not token:
        raise HTTPException(status_code=401, detail="Not authenticated")

    try:
        token = token.split(" ")[1]  # Убираем "Bearer " из заголовка
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username = payload.get("sub")
        print("Декодированный payload:", payload)  # Логирование payload

        if not username:
            raise HTTPException(status_code=401, detail="Invalid token")

        user = get_user_by_username(username)
        if user is None:
            raise HTTPException(status_code=404, detail="User not found")

        return user
    except jwt.JWTError as e:
        print("Ошибка при декодировании токена:", e)  # Логирование ошибки
        raise HTTPException(status_code=401, detail="Invalid token")


# ---------------------------------------------------------------------------------------------------------------------------------#


# Классы
# ---------------------------------------------------------------------------------------------------------------------------------#
class LoginRequest(BaseModel):
    username: str
    password: str


class RegisterRequest(BaseModel):
    username: str
    password: str
    email: str


class ProfileUpdate(BaseModel):
    email: str
    password: str


class Review(BaseModel):
    rating: int
    comment: Optional[str] = None


class Service(BaseModel):
    service_name: str
    description: str = None
    price: float = None
    duration: str = None


class CreateUserRequest(BaseModel):
    username: str
    password: str
    email: Optional[str] = None
    roleid: int


class UpdateUserRequest(BaseModel):
    username: str
    password: Optional[str] = None
    email: Optional[str] = None
    roleid: int


# Модель запроса для создания услуги
class CreateServiceRequest(BaseModel):
    name: str
    description: str = None  # Описание услуги не обязательное
    price: float
    duration: str  # Для удобства хранится как строка (например, '01:00:00')


class ReviewRequest(BaseModel):
    rating: int
    comment: str
    review_date: Optional[date] = None  # Сделаем review_date необязательным


class ReviewResponse(ReviewRequest):
    review_id: int
    service_name: str


class OrderResponse(BaseModel):
    order_id: int  # Уникальный идентификатор заказа
    order_date: date  # Дата заказа
    status: str  # Статус заказа
    first_name: Optional[str]  # Имя клиента
    last_name: Optional[str]  # Фамилия клиента
    address: Optional[str]  # Адрес клиента
    phone_number: Optional[str]  # Номер телефона клиента
    services: List[str]  # Список названий услуг


class OrderStatusUpdateRequest(BaseModel):
    status: str  # Новый статус заказа


# ---------------------------------------------------------------------------------------------------------------------------------#


# ---------------------------------------------------------------------------------------------------------------------------------#
# Домащнаяя страница
@app.get("/", response_class=HTMLResponse)
async def read_root(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


# ---------------------------------------------------------------------------------------------------------------------------------#


# ---------------------------------------------------------------------------------------------------------------------------------#
# Страница логина
@app.get("/login", response_class=HTMLResponse)
async def login(request: Request):
    return templates.TemplateResponse("login.html", {"request": request})


# Обработчик логина
@app.post("/api/login")
async def login(request: LoginRequest):
    # Получаем пользователя по имени
    user = get_user_by_username(request.username)
    roleid = user[4]
    # Проверяем, найден ли пользователь и совпадает ли пароль
    if user and check_password(
        request.password, user[2]
    ):  # user[2] - это поле пароля в базе данных
        # Создаем токен с данными пользователя (например, с его именем)

        access_token = create_access_token(data={"sub": request.username})

        return {"access_token": access_token, "token_type": "bearer", "roleid": roleid}
    else:
        raise HTTPException(status_code=401, detail="Логин или пароль введены неверно")


# Страница регистрации
@app.get("/register", response_class=HTMLResponse)
async def register(request: Request):
    return templates.TemplateResponse("register.html", {"request": request})


# Обработчик регистрации
@app.post("/api/register")
async def register(request: RegisterRequest):
    # Проверка, существует ли уже пользователь с таким именем
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM Users WHERE Username = %s;", (request.username,))
    existing_user = cursor.fetchone()
    cursor.close()
    close_db_connection(conn)

    if existing_user:
        raise HTTPException(
            status_code=400, detail="Пользователь с таким именем уже существует"
        )

    # Хэширование пароля перед сохранением
    hashed_password = hash_password(request.password)

    # Создание нового пользователя с помощью функции create_user
    try:
        user_id = create_user(
            username=request.username,
            password=hashed_password,
            email=request.email,  # Используем email из запроса
            # role_id=request.role_id,
            role_id=1,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail="Internal Server Error")

    return {"message": "User registered successfully", "user_id": user_id}


# ---------------------------------------------------------------------------------------------------------------------------------#


# ---------------------------------------------------------------------------------------------------------------------------------#
# Страница профиля
@app.get("/profile", response_class=HTMLResponse)
async def register(request: Request):
    return templates.TemplateResponse("profile.html", {"request": request})


# Обработчик профиля
@app.post("/api/profile")
def update_profile(
    profile_update: ProfileUpdate, current_user: dict = Depends(get_current_user)
):
    # Проверка, что current_user возвращается как кортеж и преобразование его в словарь
    if current_user is None or not isinstance(current_user, tuple):
        raise HTTPException(status_code=404, detail="User not found")

    # Предположим, что кортеж возвращает данные в порядке колонок: UserID, Username, Password, Email, RoleID
    columns = ["UserID", "Username", "Password", "Email", "RoleID"]
    current_user_dict = dict(zip(columns, current_user))

    user_id = current_user_dict["UserID"]

    # Получаем данные для обновления
    new_email = profile_update.email
    new_password = profile_update.password

    if new_password:
        new_password = hash_password(new_password)

    # Обновляем данные в базе данных
    conn = get_db_connection()
    cursor = conn.cursor()

    if new_email:
        cursor.execute(
            "UPDATE Users SET Email = %s WHERE UserID = %s;", (new_email, user_id)
        )

    if new_password:
        cursor.execute(
            "UPDATE Users SET Password = %s WHERE UserID = %s;", (new_password, user_id)
        )

    conn.commit()
    cursor.close()
    close_db_connection(conn)

    return {"message": "Profile updated successfully"}


# Текущие имя пользователя и email
@app.get("/api/user-info")
def get_user_info(current_user: dict = Depends(get_current_user)):
    if current_user is None or not isinstance(current_user, tuple):
        raise HTTPException(status_code=404, detail="User not found")

    # Преобразование current_user в словарь
    columns = ["UserID", "Username", "Password", "Email", "RoleID"]
    current_user_dict = dict(zip(columns, current_user))

    return {
        "username": current_user_dict["Username"],
        "email": current_user_dict["Email"],
    }


# ---------------------------------------------------------------------------------------------------------------------------------#

"""# Защищенный эндпоинт
@app.get("/users/me")
def read_users_me(current_user: dict = Depends(get_current_user)):
    return current_user"""


# ---------------------------------------------------------------------------------------------------------------------------------#
# Обработчик услуг
@app.get("/api/services")
async def get_services():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # Основной запрос для услуг
        cursor.execute(
            """
            SELECT s.ServiceID, s.ServiceName, s.Description, s.Price, s.Duration
            FROM Services s
            """
        )
        services = cursor.fetchall()

        services_list = []

        for row in services:
            service_id = row[0]
            # Получение комментариев для каждой услуги
            cursor.execute(
                """
                SELECT r.Comment, r.Rating, u.Username
                FROM Reviews r
                JOIN Clients c ON r.ClientID = c.ClientID
                JOIN Users u ON c.UserID = u.UserID
                WHERE r.ServiceID = %s;
                """,
                (service_id,),
            )
            reviews = cursor.fetchall()

            # Форматируем комментарии
            reviews_list = [
                {"comment": review[0], "rating": review[1], "username": review[2]}
                for review in reviews
            ]

            # Формируем ответ для услуги
            services_list.append(
                {
                    "id": row[0],
                    "name": row[1],
                    "description": row[2],
                    "price": float(row[3]) if row[3] else None,
                    "duration": str(row[4]) if row[4] else None,
                    "reviews": reviews_list,  # Добавляем комментарии с рейтингом
                }
            )

        cursor.close()
        conn.close()
        return services_list
    except Exception as e:
        print(f"Ошибка: {e}")
        return {"error": str(e)}


# ---------------------------------------------------------------------------------------------------------------------------------#


# ---------------------------------------------------------------------------------------------------------------------------------#
# Обработчик заказов
@app.post("/api/orders")
async def create_order(request: Request):
    current_user = get_current_user(request)

    if current_user is None:
        raise HTTPException(status_code=401, detail="Пользователь не авторизован")

    user_id = current_user[
        0
    ]  # Предполагается, что ID пользователя находится в первой позиции кортежа

    conn = None  # Инициализация переменной conn
    try:
        data = await request.json()
        print(f"Полученные данные: {data}")  # Выводим все полученные данные для отладки

        first_name = data["first_name"]
        last_name = data.get("last_name", None)
        phone = data["phone"]
        address = data["address"]
        service_ids = data.get(
            "serviceid", []
        )  # Используем "serviceid" вместо "service_ids"

        print(
            f"Список идентификаторов услуг: {service_ids}"
        )  # Проверка, что получаем service_ids

        if not service_ids:
            raise HTTPException(
                status_code=400, detail="Список идентификаторов услуг пуст"
            )

        conn = get_db_connection()
        with conn.cursor() as cursor:
            # Проверка на существующего клиента
            cursor.execute(
                """
                SELECT ClientID
                FROM Clients
                WHERE FirstName = %s AND LastName = %s AND Phone = %s AND UserID = %s;
                """,
                (first_name, last_name, phone, user_id),
            )
            result = cursor.fetchone()

            if result:
                # Клиент существует, используем его ClientID
                client_id = result[0]
            else:
                # Вставляем нового клиента
                cursor.execute(
                    """
                    INSERT INTO Clients (FirstName, LastName, Phone, Address, UserID)
                    VALUES (%s, %s, %s, %s, %s)
                    RETURNING ClientID;
                    """,
                    (first_name, last_name, phone, address, user_id),
                )
                result = cursor.fetchone()
                if result is None:
                    raise HTTPException(
                        status_code=500, detail="Не удалось получить ClientID"
                    )
                client_id = result[0]  # Новый ClientID

            # Вставка заказа с ClientID
            cursor.execute(
                """
                INSERT INTO Orders (OrderDate, Status, ClientID)
                VALUES (%s, %s, %s)
                RETURNING OrderID;
                """,
                (date.today(), "Новый", client_id),
            )
            result = cursor.fetchone()
            if result is None:
                raise HTTPException(
                    status_code=500, detail="Не удалось получить OrderID"
                )
            order_id = result[0]  # Новый OrderID

            # Вставка связей в таблицу OrderServices
            if service_ids:
                for service_id in service_ids:
                    print(
                        f"Добавляем связь: OrderID {order_id}, ServiceID {service_id}"
                    )
                    cursor.execute(
                        """
                        INSERT INTO OrderServices (OrderID, ServiceID)
                        VALUES (%s, %s);
                        """,
                        (order_id, service_id),
                    )

            conn.commit()

        return {"message": "Заказ успешно создан"}
    except Exception as e:
        if conn:
            conn.rollback()
        print(f"Ошибка при создании заказа: {e}")
        raise HTTPException(status_code=500, detail="Ошибка при создании заказа")
    finally:
        if conn:
            conn.close()


@app.get("/api/orders-info")
def get_user_orders(current_user: dict = Depends(get_current_user)):
    if not current_user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    # Получение UserID текущего пользователя
    user_id = current_user[0]  # Предположим, UserID — это первый элемент current_user

    conn = get_db_connection()
    cursor = conn.cursor()

    cursor.execute(
        """
        SELECT 
            c.FirstName, c.LastName, c.Phone, c.Address, 
            s.ServiceName, s.Price, 
            o.OrderDate, o.Status, 
            o.OrderID
        FROM Orders o
        JOIN Clients c ON o.ClientID = c.ClientID
        JOIN OrderServices os ON o.OrderID = os.OrderID
        JOIN Services s ON os.ServiceID = s.ServiceID
        WHERE c.UserID = %s;
        """,
        (user_id,),
    )

    orders = cursor.fetchall()
    columns = [
        "FirstName",
        "LastName",
        "Phone",
        "Address",
        "ServiceName",
        "Price",
        "OrderDate",
        "Status",
        "OrderID",
    ]
    orders_list = [dict(zip(columns, order)) for order in orders]

    cursor.close()
    close_db_connection(conn)

    return {"orders": orders_list}


@app.delete("/api/orders/{order_id}")
def delete_order(order_id: int, current_user: dict = Depends(get_current_user)):
    if not current_user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    conn = get_db_connection()
    cursor = conn.cursor()

    # Удаляем запись из OrderServices
    cursor.execute("DELETE FROM OrderServices WHERE OrderID = %s;", (order_id,))

    # Удаляем запись из Orders
    cursor.execute("DELETE FROM Orders WHERE OrderID = %s;", (order_id,))

    conn.commit()
    cursor.close()
    close_db_connection(conn)

    return {"message": "Order deleted successfully"}


@app.post("/api/orders/{order_id}/review")
async def add_review(order_id: int, review: Review, client=Depends(get_current_user)):
    conn = get_db_connection()
    cursor = conn.cursor()

    try:
        # Извлечение ClientID из кортежа
        client_id = client[
            0
        ]  # Предполагается, что ClientID находится на первом месте в кортеже

        if not client_id:
            raise HTTPException(status_code=400, detail="ClientID is missing")

        # Получение ServiceID через таблицу OrderServices
        cursor.execute(
            """
            SELECT os.ServiceID 
            FROM OrderServices os
            JOIN Orders o ON os.OrderID = o.OrderID
            WHERE o.OrderID = %s
            """,
            (order_id,),
        )
        result = cursor.fetchone()

        if not result:
            raise HTTPException(
                status_code=404, detail="Услуга для указанного заказа не найдена."
            )

        # Извлечение ServiceID
        service_id = result[0]

        # Добавление отзыва
        cursor.execute(
            """
            INSERT INTO Reviews (Rating, Comment, ReviewDate, ClientID, ServiceID)
            VALUES (%s, %s, %s, %s, %s)
            """,
            (review.rating, review.comment, date.today(), client_id, service_id),
        )
        conn.commit()
        return {"message": "Комментарий успешно добавлен!"}

    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))

    finally:
        cursor.close()
        conn.close()


# ---------------------------------------------------------------------------------------------------------------------------------#


# ---------------------------------------------------------------------------------------------------------------------------------#
# Обработчик Сотрудников
@app.get("/api/employees")
async def get_employees():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute(
            """
            SELECT e.EmployeeID, e.FirstName, e.LastName, e.Position, e.Email, e.Phone, e.Photo
            FROM Employees e
            """
        )
        employees = cursor.fetchall()

        employees_list = [
            {
                "id": row[0],
                "first_name": row[1],
                "last_name": row[2],
                "position": row[3],
                "email": row[4],
                "phone": row[5],
                # "photo": row[6] if row[6] else "static/images/default-avatar.jpg",
                "photo": f"/static/images/{row[6]}",  # Здесь row[6] — это имя файла, например, 'photo1.jpg'
            }
            for row in employees
        ]

        cursor.close()
        conn.close()
        return employees_list  # Возвращаем список сотрудников
    except Exception as e:
        print(f"Ошибка: {e}")
        return {"error": str(e)}


# ---------------------------------------------------------------------------------------------------------------------------------#


# ---------------------------------------------------------------------------------------------------------------------------------#
# Страница администратора
@app.get("/admin", response_class=HTMLResponse)
async def admin(request: Request):
    return templates.TemplateResponse("admin.html", {"request": request})


# ---------------------------------------------------------------------------------------------------------------------------------#


# ---------------------------------------------------------------------------------------------------------------------------------#
# Администраторская часть Users
@app.get("/api/users")
async def get_users():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT UserID, Username, Email, RoleID FROM Users;")
        users = cursor.fetchall()
        cursor.close()

        users_list = [
            {"user_id": row[0], "username": row[1], "email": row[2], "roleid": row[3]}
            for row in users
        ]
        conn.close()

        return users_list
    except Exception as e:
        raise HTTPException(status_code=500, detail="Internal Server Error")


@app.post("/api/users")
async def create_user_by_admin(request: CreateUserRequest):
    try:
        # Проверка, существует ли уже пользователь с таким именем
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM Users WHERE Username = %s;", (request.username,))
        existing_user = cursor.fetchone()
        cursor.close()

        if existing_user:
            raise HTTPException(
                status_code=400, detail="Пользователь с таким именем уже существует"
            )

        # Хэширование пароля перед сохранением
        hashed_password = hash_password(request.password)

        # Создание нового пользователя
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO Users (Username, Password, Email, RoleID) VALUES (%s, %s, %s, %s) RETURNING UserID;",
            (request.username, hashed_password, request.email, request.roleid),
        )
        user_id = cursor.fetchone()[0]
        conn.commit()
        cursor.close()
        conn.close()

        return {"message": "User created successfully", "user_id": user_id}
    except Exception as e:
        raise HTTPException(status_code=500, detail="Internal Server Error")


@app.delete("/api/users/{username}")
async def delete_user(username: str):
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM Users WHERE UserName = %s;", (username,))
        user = cursor.fetchone()
        cursor.close()

        if not user:
            raise HTTPException(status_code=404, detail="User not found")

        cursor = conn.cursor()
        cursor.execute("DELETE FROM Users WHERE UserName = %s;", (username,))
        conn.commit()
        cursor.close()
        conn.close()

        return {"message": "User deleted successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail="Internal Server Error")


@app.put("/api/users/{username}")
async def update_user(username: str, request: UpdateUserRequest):
    try:
        # Проверяем, существует ли пользователь с таким username
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM Users WHERE Username = %s;", (username,))
        user = cursor.fetchone()
        cursor.close()

        if not user:
            raise HTTPException(status_code=404, detail="User not found")

        # Подготовим данные для обновления
        updated_data = {
            "username": request.username,
            "email": request.email,
            "roleid": request.roleid,
        }

        # Если есть новый пароль, хешируем его
        if request.password:
            updated_data["password"] = hash_password(request.password)
        else:
            updated_data["password"] = hash_password(
                user[1]
            )  # Если пароля нет в запросе, оставляем старый пароль

        # Строим запрос для обновления
        update_query = """
        UPDATE Users 
        SET Username = %s, Password = %s, Email = %s, RoleID = %s 
        WHERE Username = %s;
        """

        # Выполняем обновление в базе данных
        cursor = conn.cursor()
        cursor.execute(
            update_query,
            (
                updated_data["username"],
                updated_data["password"],
                updated_data["email"]
                or user[2],  # Если email не передан, оставляем старое значение
                updated_data["roleid"],
                username,  # Обновление по username
            ),
        )
        conn.commit()
        cursor.close()
        conn.close()

        return {"message": "User updated successfully"}
    except Exception as e:
        print(f"Error during update: {e}")  # Для отладки
        raise HTTPException(status_code=500, detail="Internal Server Error")


# ---------------------------------------------------------------------------------------------------------------------------------#
# Администраторская часть Services


# Эндпоинт для получения списка всех услуг
@app.get("/api/services")
async def get_services():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # Запрос для получения всех услуг
        cursor.execute(
            "SELECT ServiceID, ServiceName, Description, Price, Duration FROM Services;"
        )
        services = cursor.fetchall()

        # Закрытие курсора
        cursor.close()

        # Формируем список услуг
        services_list = [
            {
                "ServiceID": row[0],
                "ServiceName": row[1],
                "Description": row[2] if row[2] else "Нет описания",
                "Price": row[3] if row[3] else "Не указано",
                "Duration": row[4] if row[4] else "Не указано",
            }
            for row in services
        ]

        # Закрытие соединения
        conn.close()

        return services_list

    except Exception as e:
        raise HTTPException(status_code=500, detail="Internal Server Error")


@app.post("/api/services")
async def create_service_by_admin(request: CreateServiceRequest):
    try:
        # Получаем соединение с БД
        conn = get_db_connection()
        cursor = conn.cursor()

        # Проверка, существует ли услуга с таким названием
        cursor.execute(
            "SELECT * FROM Services WHERE ServiceName = %s;", (request.name,)
        )
        existing_service = cursor.fetchone()

        if existing_service:
            cursor.close()
            conn.close()
            raise HTTPException(status_code=400, detail="Service name already taken")

        # Добавляем новую услугу
        cursor.execute(
            "INSERT INTO Services (ServiceName, Description, Price, Duration) VALUES (%s, %s, %s, %s) RETURNING ServiceID;",
            (request.name, request.description, request.price, request.duration),
        )
        service_id = cursor.fetchone()[0]
        conn.commit()

        cursor.close()
        conn.close()

        return {"message": "Service created successfully", "service_id": service_id}

    except Exception as e:
        print(f"Error while creating service: {e}")
        raise HTTPException(status_code=500, detail="Internal Server Error")


@app.put("/api/services/{service_id}")
async def update_service(service_id: int, request: CreateServiceRequest):
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # Проверка, существует ли услуга с таким названием
        cursor.execute(
            "SELECT * FROM Services WHERE ServiceName = %s AND ServiceID != %s;",
            (request.name, service_id),
        )
        existing_service = cursor.fetchone()

        if existing_service:
            cursor.close()
            conn.close()
            raise HTTPException(status_code=400, detail="Service name already taken")

        # Обновляем информацию об услуге
        cursor.execute(
            "UPDATE Services SET ServiceName = %s, Description = %s, Price = %s, Duration = %s WHERE ServiceID = %s;",
            (
                request.name,
                request.description,
                request.price,
                request.duration,
                service_id,
            ),
        )
        conn.commit()

        cursor.close()
        conn.close()

        return {"message": "Service updated successfully"}

    except Exception as e:
        raise HTTPException(status_code=500, detail="Internal Server Error")


@app.delete("/api/services/{service_id}")
async def delete_service(service_id: int):
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # Удаляем услугу
        cursor.execute("DELETE FROM Services WHERE ServiceID = %s;", (service_id,))
        conn.commit()

        cursor.close()
        conn.close()

        return {"message": "Service deleted successfully"}

    except Exception as e:
        raise HTTPException(status_code=500, detail="Internal Server Error")


# ---------------------------------------------------------------------------------------------------------------------------------#
# Администраторская часть Reviews


# Эндпоинт для получения всех отзывов
@app.get("/api/reviews", response_model=List[ReviewResponse])
async def get_reviews():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute(
            """
            SELECT r.ReviewID, r.Rating, r.Comment, r.ReviewDate, s.ServiceName
            FROM Reviews r
            JOIN Services s ON r.ServiceID = s.ServiceID;
        """
        )
        reviews = cursor.fetchall()
        cursor.close()
        conn.close()

        return [
            {
                "review_id": review[0],
                "rating": review[1],
                "comment": review[2],
                "review_date": review[3],
                "service_name": review[4],
            }
            for review in reviews
        ]
    except Exception as e:
        raise HTTPException(status_code=500, detail="Internal Server Error")


@app.put("/api/reviews/{review_id}", response_model=ReviewResponse)
async def update_review(review_id: int, review_request: ReviewRequest):
    try:
        print("Received data:", review_request)  # Логируем полученные данные
        conn = get_db_connection()
        cursor = conn.cursor()

        # Обновляем отзыв в базе данных
        cursor.execute(
            """
            UPDATE Reviews
            SET Rating = %s, Comment = %s
            WHERE ReviewID = %s
            """,
            (review_request.rating, review_request.comment, review_id),
        )

        conn.commit()

        # Получаем обновленные данные для возвращения
        cursor.execute(
            """
            SELECT r.ReviewID, r.Rating, r.Comment, r.ReviewDate, s.ServiceName
            FROM Reviews r
            JOIN Services s ON r.ServiceID = s.ServiceID
            WHERE r.ReviewID = %s;
            """,
            (review_id,),
        )
        updated_review = cursor.fetchone()

        cursor.close()
        conn.close()

        if not updated_review:
            raise HTTPException(status_code=404, detail="Review not found")

        return {
            "review_id": updated_review[0],
            "rating": updated_review[1],
            "comment": updated_review[2],
            "review_date": updated_review[3],
            "service_name": updated_review[4],
        }

    except Exception as e:
        print("Error:", e)  # Логируем ошибку
        raise HTTPException(status_code=500, detail="Internal Server Error")


# Эндпоинт для удаления отзыва
@app.delete("/api/reviews/{review_id}")
async def delete_review(review_id: int):
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        # Удаляем отзыв из базы данных
        cursor.execute(
            """
            DELETE FROM Reviews
            WHERE ReviewID = %s
            """,
            (review_id,),
        )

        conn.commit()

        cursor.close()
        conn.close()

        return {"message": f"Review with ID {review_id} deleted successfully."}

    except Exception as e:
        print(e)  # Логирование ошибки (для отладки)
        raise HTTPException(status_code=500, detail="Internal Server Error")


# ---------------------------------------------------------------------------------------------------------------------------------#
# Администраторская часть Orders


@app.get("/api/ordersadmin", response_model=List[OrderResponse])
async def get_orders():
    try:
        # Подключаемся к базе данных
        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute(
            """
            SELECT 
            o.OrderID, 
            o.OrderDate, 
            o.Status, 
            c.FirstName, 
            c.LastName, 
            c.Address,  -- Адрес клиента
            c.Phone,  -- Номер телефона клиента
            ARRAY_AGG(s.ServiceName) AS Services
            FROM Orders o
            LEFT JOIN Clients c ON o.ClientID = c.ClientID
            LEFT JOIN OrderServices os ON o.OrderID = os.OrderID
            LEFT JOIN Services s ON os.ServiceID = s.ServiceID
            GROUP BY o.OrderID, o.OrderDate, o.Status, c.FirstName, c.LastName, c.Address, c.Phone;
            """
        )
        orders = cursor.fetchall()

        cursor.close()
        conn.close()

        return [
            {
                "order_id": order[0],
                "order_date": order[1],
                "status": order[2],
                "first_name": order[3],  # Имя клиента
                "last_name": order[4],  # Фамилия клиента
                "address": order[5],  # Адрес клиента
                "phone_number": order[6],  # Номер телефона клиента
                "services": order[7],  # Список услуг
            }
            for order in orders
        ]
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Internal Server Error: {str(e)}")


@app.delete("/api/ordersadmin/{order_id}", status_code=204)
async def delete_order(order_id: int):
    try:
        # Подключаемся к базе данных
        conn = get_db_connection()
        cursor = conn.cursor()

        # Удаляем записи из таблицы OrderServices
        cursor.execute("DELETE FROM OrderServices WHERE OrderID = %s;", (order_id,))

        # Удаляем сам заказ
        cursor.execute("DELETE FROM Orders WHERE OrderID = %s;", (order_id,))

        # Фиксируем изменения
        conn.commit()

        cursor.close()
        conn.close()

        return Response(status_code=204)  # Успешное удаление, без содержимого
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Internal Server Error: {str(e)}")


@app.put("/api/ordersadmin/{order_id}", status_code=200)
async def update_order_status(order_id: int, order_update: OrderStatusUpdateRequest):
    try:
        # Подключаемся к базе данных
        conn = get_db_connection()
        cursor = conn.cursor()

        # Обновляем только статус заказа
        cursor.execute(
            """
            UPDATE Orders
            SET Status = %s
            WHERE OrderID = %s;
            """,
            (order_update.status, order_id),
        )

        # Фиксируем изменения
        conn.commit()

        cursor.close()
        conn.close()

        return {"message": "Статус заказа успешно обновлен"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Internal Server Error: {str(e)}")


# ---------------------------------------------------------------------------------------------------------------------------------#
# Администраторская часть Employees
# Получение всех сотрудников с расширенной информацией
# Pydantic models for request and response validation
class EmployeeBase(BaseModel):
    first_name: str
    last_name: str
    position: str = None
    hire_date: str = None
    email: str = None
    phone: str = None
    photo: str = None


class EmployeeResponse(EmployeeBase):
    employee_id: int


class EmployeeUpdate(EmployeeBase):
    pass


# ---------------------------------------------------------------------------------------------------------------------------------#
# Get all employees
@app.get("/api/employeesadmin", response_model=List[EmployeeResponse])
async def get_employees():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute("SELECT * FROM Employees;")
        employees = cursor.fetchall()

        cursor.close()
        conn.close()

        return [
            {
                "employee_id": employee[0],
                "first_name": employee[1],
                "last_name": employee[2],
                "position": employee[3],
                "hire_date": (
                    employee[4].strftime("%Y-%m-%d") if employee[4] else None
                ),  # Convert date to string
                "email": employee[5],
                "phone": employee[6],
                "photo": employee[7],
            }
            for employee in employees
        ]
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Internal Server Error: {str(e)}")


# ---------------------------------------------------------------------------------------------------------------------------------#
# Create new employee
@app.post("/api/employeesadmin", response_model=EmployeeResponse)
async def create_employee(employee: EmployeeBase):
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute(
            """
            INSERT INTO Employees (FirstName, LastName, Position, HireDate, Email, Phone, Photo)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
            RETURNING EmployeeID;
            """,
            (
                employee.first_name,
                employee.last_name,
                employee.position,
                employee.hire_date,
                employee.email,
                employee.phone,
                employee.photo,
            ),
        )
        employee_id = cursor.fetchone()[0]
        conn.commit()

        cursor.close()
        conn.close()

        return {**employee.dict(), "employee_id": employee_id}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Internal Server Error: {str(e)}")


# ---------------------------------------------------------------------------------------------------------------------------------#
# Update an employee
@app.put("/api/employeesadmin/{employee_id}", response_model=EmployeeResponse)
async def update_employee(employee_id: int, employee_update: EmployeeUpdate):
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute(
            """
            UPDATE Employees
            SET FirstName = %s, LastName = %s, Position = %s, HireDate = %s, Email = %s, Phone = %s, Photo = %s
            WHERE EmployeeID = %s;
            """,
            (
                employee_update.first_name,
                employee_update.last_name,
                employee_update.position,
                employee_update.hire_date,
                employee_update.email,
                employee_update.phone,
                employee_update.photo,
                employee_id,
            ),
        )
        conn.commit()

        cursor.close()
        conn.close()

        return {**employee_update.dict(), "employee_id": employee_id}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Internal Server Error: {str(e)}")


# ---------------------------------------------------------------------------------------------------------------------------------#
# Delete an employee
@app.delete("/api/employeesadmin/{employee_id}", status_code=204)
async def delete_employee(employee_id: int):
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute("DELETE FROM Employees WHERE EmployeeID = %s;", (employee_id,))
        conn.commit()

        cursor.close()
        conn.close()

        return Response(status_code=204)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Internal Server Error: {str(e)}")


# ---------------------------------------------------------------------------------------------------------------------------------#
# Semd Email
from fastapi import FastAPI, HTTPException, BackgroundTasks
from pydantic import BaseModel, EmailStr
from typing import List
import psycopg2
from fpdf import FPDF
from openpyxl import Workbook
from docx import Document
import aiosmtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders

SMTP_CONFIG = {
    "host": "smtp.gmail.com",
    "port": 587,
    "username": "your_email",
    "password": "needed_password",
}


class SendServicesRequest(BaseModel):
    email: EmailStr
    format: str


# Поддерживаемые форматы
SUPPORTED_FORMATS = ["pdf", "word", "excel"]


# Функция для извлечения услуг из базы данных
def fetch_services_from_db():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT ServiceName, Description, Price, Duration FROM Services")
        services = cursor.fetchall()
        cursor.close()
        conn.close()

        # Преобразуем данные в читаемый формат
        return [
            {
                "name": row[0],
                "description": row[1],
                "price": row[2],
                "duration": str(row[3]) if row[3] else "Не указана",
            }
            for row in services
        ]
    except Exception as e:
        print(f"Ошибка при извлечении данных из БД: {e}")
        return []


# Функция генерации документа
from fpdf import FPDF


def generate_document(services, format):
    filename = f"services.{format}"
    # Присваиваем правильное расширение в зависимости от формата
    if format == "pdf":
        pdf = FPDF()
        pdf.add_page()

        # Замените путь на путь к вашему TTF файлу
        pdf.add_font("Arial", "", "C:\\Users\\Yuriy\\Downloads\\Arial.ttf", uni=True)
        pdf.set_font("Arial", size=12)

        pdf.cell(200, 10, txt="Услуги", ln=True, align="C")

        for service in services:
            pdf.cell(
                200,
                10,
                txt=f"{service['name']} - {service['description']} - {service['price']} руб - {service['duration']}",
                ln=True,
            )

        pdf.output(filename)

    elif format == "excel":
        filename = "services.xlsx"

        # Генерация Excel
        wb = Workbook()
        ws = wb.active
        ws.title = "Услуги"
        ws.append(["Название", "Описание", "Цена (руб)", "Длительность"])
        for service in services:
            ws.append(
                [
                    service["name"],
                    service["description"],
                    service["price"],
                    service["duration"],
                ]
            )
        wb.save(filename)

    elif format == "word":
        filename = "services.docx"

        # Генерация Word
        doc = Document()
        doc.add_heading("Услуги", level=1)
        for service in services:
            doc.add_paragraph(
                f"{service['name']} - {service['description']} - {service['price']} руб - {service['duration']}"
            )
        doc.save(filename)

    return filename


# Функция отправки email
async def send_email(email: str, filename: str):
    message = MIMEMultipart()
    message["From"] = SMTP_CONFIG["username"]
    message["To"] = email
    message["Subject"] = "Услуги"

    # Тело письма
    message.attach(MIMEBase("text", "plain"))
    message.attach(MIMEBase("text", "plain"))

    with open(filename, "rb") as attachment:
        part = MIMEBase("application", "octet-stream")
        part.set_payload(attachment.read())
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", f"attachment; filename={filename}")
        message.attach(part)

    try:
        await aiosmtplib.send(
            message,
            hostname=SMTP_CONFIG["host"],
            port=SMTP_CONFIG["port"],
            username=SMTP_CONFIG["username"],
            password=SMTP_CONFIG["password"],
            start_tls=True,
        )
    except Exception as e:
        print(f"Ошибка при отправке email: {e}")
        raise HTTPException(status_code=500, detail="Ошибка при отправке email")


# Эндпоинт для отправки услуг
@app.post("/api/send-services")
async def send_services(
    request: SendServicesRequest, background_tasks: BackgroundTasks
):
    # Проверяем формат
    if request.format not in SUPPORTED_FORMATS:
        raise HTTPException(
            status_code=400,
            detail="Неподдерживаемый формат. Доступны: pdf, word, excel",
        )

    # Извлекаем данные из базы
    services = fetch_services_from_db()
    if not services:
        raise HTTPException(status_code=404, detail="Нет доступных услуг")

    # Генерируем документ
    filename = generate_document(services, request.format)

    # Отправляем email в фоновом процессе
    background_tasks.add_task(send_email, request.email, filename)

    return {"message": "Услуги будут отправлены на указанный email."}
